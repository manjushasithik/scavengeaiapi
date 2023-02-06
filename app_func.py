"""
Simple app to upload an image via a web form 
and view the inference results on the image in the browser.
"""
from subprocess import STDOUT, check_call , call,run

# import argparse
import io
import math
import os
from PIL import Image

import numpy as np
from base64 import b64encode
# import pythoncom
import re
from docx.enum.table import WD_ALIGN_VERTICAL

import base64
import torch


import pandas as pd

import openpyxl
from flask import Flask, render_template, request, redirect,jsonify,send_file

import os
import json
# from docx2pdf import convert
from docx import Document # for pdf format
from docx.shared import Pt # for pdf format
from docx.shared import Inches


def predict(data_raw):

#     github='ultralytics/yolov5'
#     torch.hub.list(github, trust_repo=True)
#     model = torch.hub.load("ultralytics/yolov5", "custom", path = "./rings18.pt", force_reload=True)
    
#     model.classes=[3 ,10,11 ,12, 17]


    #print("Here")
    data_raw = data_raw
    # print("data_raw:",data_raw)
    # print("***********************************")
    # data=dict(itertools.islice(data_raw.items(), 8,len(data_raw)))
    # data_id1 =dict(itertools.islice(data_raw.items(), 4,8))
    # data_id=dict(itertools.islice(data_raw.items(), 0,4))
    
    l_raw=list(data_raw.keys())
    # print(l_raw)
    l_id=['VESSEL_OBJECT_ID','EQUIPMENT_CODE', 'EQUIPMENT_ID', 'JOB_PLAN_ID','JOB_ID','LOG_ID']
    l_pdf_para=['Vessel','Hull_No','Vessel_Type','Local_Start_Time','Time_Zone1','Local_End_Time','Time_Zone2','Form_No','IMO_No','Maker','Model','License_Builder','Serial_No','MCR','Speed_at_MCR','Bore','Stroke','Maker_T','Model_T','Total_Running_Hour','Cylinder_Oil_type','Normal_service_load_in_percentage_of_MCR','Scrubber','Position','Cylinder_Oil_feed_rate','Inspected_by_Rank','Fuel_Sulphur_percentage']
    l_data=[]
    for i in l_raw:
        if i not in l_id and i not in l_pdf_para:
            l_data.append(i)
    data_id={}
    for i in l_id:
        if i not in data_id.keys():
            data_id[i]=data_raw[i]
    data_pdf_para={}
    for i in l_pdf_para:
        if i not in data_pdf_para.keys():
            data_pdf_para[i]=data_raw[i]
    data={}
    for i in l_data:
        if i not in data.keys():
            data[i]=data_raw[i]
    data_keys=list(data.keys())
    # print("data:",data)
    # print("data_keys:",data_keys)
    # print("data_id:",data_id)   
    # print("data_pdf_para:",data_pdf_para)
    cyl_index=-1
    defect_df_all_cyl={}
    img_list=[]
    obj_lst=[]
    cyl_obj={}
    data = {k: v for k, v in data.items() if v is not None}
    # print("data:",data.keys())
    data_k=list(data.keys())
    cyl_num_temp=[re.split(r'(\d+)', s) for s in data_k]
    cyl_num=[i[1] for i in cyl_num_temp]
    len_tensor=0
    cyl_pred={}

    for (cyl_number,cyl) in  zip(data.keys(),data.values()):
        cyl_index=cyl_index+1
#---
        header, encoded = cyl.split(",", 1)
        data = base64.b64decode(encoded)

        with open("image.png", "wb") as f:
            f.write(data)   

        img = Image.open(io.BytesIO(data))
        results = model(img, size=640)
        img = np.squeeze(results.render())
        # datatoexcel = pd.ExcelWriter('results.xlsx')
        # results.to_excel(datatoexcel)
        # datatoexcel.save()

        #print("RESULT=======",results)
        file_object = io.BytesIO()
        
        data = Image.fromarray(img)
        data.save(file_object, 'JPEG')
        
        base64img = "data:image/png;base64,"+b64encode(file_object.getvalue()).decode('ascii')

        res_tensor=results.xyxy[0]  # im1 predictions (tensor)
        # print("res_tensor:",res_tensor)
        len_tensor=len(res_tensor)
        if len(res_tensor)==0:
            img_list.append(file_object)
            cyl_pred[cyl_number]={"Overall_fault_id":"-99","Overall_Rating":"3","Recommendation":"(Satisfactory) - No fault - No recommendation"}
            defect_df_all_cyl["cylinder"+(data_keys[cyl_index])]={'lubrication': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}, 
            'surface': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}, 
            'deposits': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}, 
            'breakage': {'Ring1': '*', 'Ring2': '*', 'Ring3': '*', 'Ring4': '*'}}
            obj_lst.append({'Fault_id':"999",'Rating':"3","Recommendation":"(Satisfactory) - No fault - No recommendation"})
                 
        else:
            h=data.height

        
    ################################################################################################################## Object #############################################################################################
            ##{"3"-C: '4',"12"-S: '3',"17"-LC: '5',"10"-O: '2', "11"-OB:'1'}
            faultdict = {"3": '4',"12": '3',"17": '5',"10": '2', "11":'1'}
            conf_lvl = {"0":"Unacceptable","1":"Marginal","2":"Fair","3":"Satisfactory"}
            
            #result coverted to Dataframe
            df = results.pandas().xyxy[0]  
            # print("dataframe:",df)

            df["rings"] = list(map(lambda x: "ring1" if x/h <=0.25 else ("ring2" if x/h <=0.45 else ("ring3" if x/h <=0.75 else "ring4")),df['ymin'])) #Map rings to rings column
            df["rating"] = list(map(lambda x: "Satisfactory" if x<0.25 else ("Fair" if x <0.5 else ("Marginal" if x<0.75 else "Unacceptable") ),df['confidence'])) #Map confidence level to conf column
            df["confidence_lvl"] = list(map(lambda x: "3" if x<0.25 else ("2" if x <0.5 else ("1" if x<0.75 else "0") ),df['confidence']))
            df["class"] = list(map(lambda x: "4" if x==3 else ("3" if x ==12 else ("5" if x==17 else ("2" if x==10 else "1")) ),df['class']))
            df["fault"] = list(map(lambda x: "Oil Black" if x=="1" else ("Too Much Oil" if x =="2" else ("Scratch" if x=="3" else ("Collapsed" if x=="4" else "Carbon")) ),df['class'])) #Map class to fault column
            # print("df:",df)
            #List of list of all parameters
            defect=[]
            for r in range(len(df)):
                ls=[]
                ls.append(df.loc[r]['rings'])
                ls.append(df.loc[r]["rating"])
                ls.append(df.loc[r]["class"])
                ls.append(df.loc[r]["fault"])
                defect.append(ls)
            # print("defect_beforeloop:",defect)
            # print(len(defect))
            #Appending recommendations dummy values to defects list
            # for k in faultdict.values():
            #     for d in range(len(defect)):
            #         if defect[d][1]=="Unacceptable" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[0]))
            #         elif defect[d][1]=="Marginal" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[1]))
            #         elif defect[d][1]=="Fair" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[2]))
            #         elif defect[d][1]=="Satisfactory" and defect[d][2]==k:   
            #             defect[d].append("s"+str(k)+ str(list(conf_lvl.keys())[3]))
            for k in faultdict.values():
                for d in range(len(defect)):
                    if defect[d][2]==k and k=="5":
                        defect[d].append("""1) Replace or Overhaul fuel injector to avoid improper combustion
                                    2) Adjust fuel temperature/viscocity to attain correct viscocity as per Maker's recommendation
                                    3) Check the condition of piston rings free movement (Gas sealing)
                                    4) Adjust the Cylinder Oil Feed rate to avoid over lubrication to avoid formation of carbon deposits
                                    """)
                    elif defect[d][2]==k and k=="4":
                        defect[d].append("""1) Replace Piston Rings
                                    2) Check for Carbon deposits in the ring groove
                                    3) Check vertical ring clearance
                                    4) Check for Partial sticking
                                    5) Check for Poor sealing between the ring and the ring groove floor.
                                    6) Check for Clover leafing 
                                    7) Check for Ring end chamfers.
                                    8) Check for too large ring edge radii.
                                    9) Check for Continual striking against wear ridges, or other irregularities in the cylinder wall.
                                    """)
                    elif defect[d][2]==k and k=="3":
                        defect[d].append("""1) Adjust the Cylinder Oil Feed rate
                                    2) Carry out Drain oil analysis (On board or send ashore)
                                    3) Carry out or land samples for Fuel oil analysis
                                    4) Check for Hard abrasive particles
                                    """)             
                    elif defect[d][2]==k and k=="1":
                        defect[d].append("""1) Check Fuel injectors for leakage
                                    2) Check for carbon deposits
                                    """)
                    elif defect[d][2]==k and k=="2":
                        defect[d].append("""1) Adjust the Cylinder Oil Feed rate
                                    2) Carry out Drain oil analysis (On board or send ashore)
                                    3) Adjust feed rate to obtain optimum residual BN
                                    """)

            # print("defect_after_loop:",defect)
            #Overall defect summary
            tensor = {}
            if df["class"].nunique() > 1:
                tensor["multi_defect"]= df.confidence.max()
            else:
                tensor[df.loc[np.argmax(df["confidence"])]['class']] = df.loc[np.argmax(df["confidence"])]['confidence']
            defect.insert(0, tensor)

            #Final object
            # str1,str2 = '',''
            # for item in tensor:
            #     str2 += '('+item + ')-(' + tensor[item] + ')'
            # for l1 in range(1,len(defect)):
            #     str1 = "("+ defect[l1][1]+') - ('+defect[l1][2]+') -'+defect[l1][3]+'-'+defect[l1][4]
            #     str2 += '||'+str1

            # obj={}
            # print("defect:",defect)
            str0,str1 = '',''
            for l1 in range(1,len(defect)):
                str0 = "("+ defect[l1][1]+') - '+defect[l1][3]+' - '+defect[l1][4] + " || " ##Note: Add ring number.
                str1 += str0
            #     print("----------------------",l1)    
            str1

            obj={}
            
            if df["class"].nunique() > 1:
                obj["Fault_id"] = 99
                obj["Rating"]= df.loc[np.argmax(df["confidence"])]['confidence_lvl']
                obj["Recommendation"] = str1
            else:
                obj["Fault_id"] = df.loc[np.argmax(df["confidence"])]['class']
                obj["Rating"] = df.loc[np.argmax(df["confidence"])]['confidence_lvl']
                obj["Recommendation"] = str1
            # print("obj:",obj)
            obj_lst.append(obj)
            cyl_pred[cyl_number]=obj_lst

    
            # defect.insert(0, tensor)



    ########################################################################################################################################################################################################################

            img_list.append(file_object) ########### adding images in docx #########################
            #print(results.pandas().xyxy[0] ) # im1 predictions (pandas)
            #print("y ",res_tensor[0][1])
            #print("c ",int(res_tensor[0][5]))
            #print("tensor len",len(res_tensor))
            
            rings=[]
            for i in range(0,len(res_tensor)):
                
                #print("percent=====",res_tensor[i][1]/h)
                if res_tensor[i][1]/h <=.25 :
                    rings.append({"1":int(res_tensor[i][5])})
                elif res_tensor[i][1]/h <=.45 :
                    rings.append({"2":int(res_tensor[i][5])})
                elif res_tensor[i][1]/h <=.75 :
                    rings.append({"3":int(res_tensor[i][5])})
                elif res_tensor[i][1]/h >.75 :
                    rings.append({"4":int(res_tensor[i][5])})
        
            def_section_brk=set()
            def_section_lub1=set()
            def_section_surf=set()
            def_section_dep=set()
            def_section_lub2=set()
            def_section_brk_ls={}
            
            def_section_lub_ls={}
            def_section_surf_ls={}
            def_section_dep_ls={}

            # Fault_id added
            
            try:
                for ring_no in range(1,5):
                    
                    def_section_lub_ls["Ring"+str(ring_no)]="*"
            except Exception :
                print("Excepetion")

            try:
                for ring_no in range(1,5):
                    
                    
                    def_section_surf_ls["Ring"+str(ring_no)]="*"
            except Exception :
                print("Excepetion")
            try:
                for ring_no in range(1,5):
                    
                    
                    def_section_dep_ls["Ring"+str(ring_no)]="*"
                    
            except Exception :
                print("Excepetion")
            try:
                for ring_no in range(1,5):
                
                    def_section_brk_ls["Ring"+str(ring_no)]="*"
                
            except Exception :
                print("Excepetion")

            for ring in rings:
                # print(ring.values())
            
                if(list(ring.values())[0]==3):# if collapsed
                    def_section_brk.add(list(ring.keys())[0]) # assign ring number
                if(list(ring.values())[0]==12): 
                    def_section_surf.add(list(ring.keys())[0])
                if(list(ring.values())[0]==11 ):
                    def_section_lub1.add(list(ring.keys())[0])
                if(list(ring.values())[0]==17):
                    def_section_dep.add(list(ring.keys())[0])
                if(list(ring.values())[0]==10):
                    def_section_lub2.add(list(ring.keys())[0])
        
            # for brk in def_section_brk :
            #     #print({"Ring"+brk:"C"})
            #     def_section_brk_ls.update({"Ring"+brk:"C"})
            # for brk in def_section_surf :

            #     def_section_surf_ls.update({"Ring"+brk:"S"})
            # for brk in def_section_dep :
            #     #print({"Ring"+brk:"LC"})
            #     def_section_dep_ls.update({"Ring"+brk:"LC"})

            # for brk in def_section_lub1 :
            #     def_section_lub_ls.update({"Ring"+brk:"OB"})
            # for brk in def_section_lub2 :
            #     #print({"Ring"+brk:"OB"})
            #     if brk not in(list(def_section_lub1)):
            #         def_section_lub_ls.update({"Ring"+brk:"O"})
            # #print(def_section_dep_ls)

            for brk in def_section_brk :
                # #faultdict = {"3"-C: '4',"12"-S: '3',"17"-LC: '5',"10"-O: '2', "11"-OB:'1'}
                #print({"Ring"+brk:faultdict["C"]})
                def_section_brk_ls.update({"Ring"+brk:faultdict["3"]})
            for brk in def_section_surf :
                #print({"Ring"+brk:faultdict["S"]})
                def_section_surf_ls.update({"Ring"+brk:faultdict["12"]})
            for brk in def_section_dep :
                #print({"Ring"+brk:faultdict["LC"]})
                def_section_dep_ls.update({"Ring"+brk:faultdict["17"]})

            for brk in def_section_lub1 :
                #print({"Ring"+brk:faultdict["OB"]})
                def_section_lub_ls.update({"Ring"+brk:faultdict["11"]})
            for brk in def_section_lub2 :
                #print({"Ring"+brk:faultdict["OB"]})
                def_section_lub_ls.update({"Ring"+brk:faultdict["11"]})
                if brk not in(list(def_section_lub1)):
                    def_section_lub_ls.update({"Ring"+brk:faultdict["10"]})
            #print(def_section_dep_ls)
        

            #print("def_section_lub_ls", def_section_lub_ls)
            #js_data=results.pandas().xyxy[0].to_json(orient="records")
            
            defect_df=  {"lubrication":def_section_lub_ls, "surface":def_section_surf_ls,"deposits":def_section_dep_ls,"breakage":def_section_brk_ls
            #, "image":base64img
            }
            # print("defect_df:",defect_df)
            # print("data_k",data_k)
            # print("cyl_index",cyl_index)
            defect_df_all_cyl["cylinder"+(data_keys[cyl_index])]=defect_df
            # print("defect_df_all_cyl",defect_df_all_cyl)
            # ##print(defect_df)
            # #print("------------------defect-------------------------------")
        # print("len of tensor:",len_tensor)
    #print(defect_df_all_cyl)

    selection_lubrication=[]
    selection_surface=[]
    selection_deposits=[]
    selection_brekage=[]
        
            
    user_data=[]
            
    cyls=defect_df_all_cyl.keys()
    print(cyls)
        
        
        # for cyl in cyls:
        #     #print(type(cyl))
        #     if cyl.startswith("cylinder") and len_tensor>0 :

        #         selection_lubrication.append(list(defect_df_all_cyl[cyl]['lubrication'].values()))
        

        #         selection_surface.append(list(defect_df_all_cyl[cyl]['surface'].values()))
        

        #         selection_brekage.append(list(defect_df_all_cyl[cyl]['breakage'].values()))
        

        #         selection_deposits.append(list(defect_df_all_cyl[cyl]['deposits'].values()))
        #         print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)
        #     # elif cyl.startswith("cylinder") and len_tensor==0:
        #     elif cyl
        #         selection_lubrication.append(["*","*","*","*"])
        #         selection_surface.append(["*","*","*","*"])
        #         selection_brekage.append(["*","*","*","*"])
        #         selection_deposits.append(["*","*","*","*"])
        #         print("**************************")
        #         print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)

    #if len_tensor>0:
    for cyl in cyls:
        #print(type(cyl))
        if cyl.startswith("cylinder"):

            selection_lubrication.append(list(defect_df_all_cyl[cyl]['lubrication'].values()))
    

            selection_surface.append(list(defect_df_all_cyl[cyl]['surface'].values()))
    

            selection_brekage.append(list(defect_df_all_cyl[cyl]['breakage'].values()))
    

            selection_deposits.append(list(defect_df_all_cyl[cyl]['deposits'].values()))
        # print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)
        # elif cyl.startswith("cylinder") and len_tensor==0:
    #else:
    # selection_lubrication.append(["*","*","*","*"])
    # selection_surface.append(["*","*","*","*"])
    # selection_brekage.append(["*","*","*","*"])
    # selection_deposits.append(["*","*","*","*"])
    # print("selection_deposits",selection_deposits,selection_brekage,selection_surface,selection_lubrication)



    pred_per_cyl_lubrication_rev = [[selection_lubrication[j][i] for j in range(len(selection_lubrication))] for i in range(len(selection_lubrication[0]))]
    pred_per_cyl_surface_rev = [[selection_surface[j][i] for j in range(len(selection_surface))] for i in range(len(selection_surface[0]))]
    pred_per_cyl_deposits_rev = [[selection_deposits[j][i] for j in range(len(selection_deposits))] for i in range(len(selection_deposits[0]))]
    pred_per_cyl_breakage_rev = [[selection_brekage[j][i] for j in range(len(selection_brekage))] for i in range(len(selection_brekage[0]))]
    #print(pred_per_cyl_lubrication_rev)
            
    
    
    # for i,j in zip(data_keys,obj_lst): ##Final Object
    #     cyl_obj[i]=j
    # final_cyl_obj={**cyl_obj, **data_id}
    # print("Object:", final_cyl_obj)
    # result_obj = json.loads(json.dumps(final_cyl_obj))  
    # print(result_obj)

############################################################### PDF ##################################################

    #print(len(data_keys))
    doc = Document('sample_report.docx')
    doc.tables #a list of all tables in document
    # # Table 0
    # doc.tables[0].cell(0, 1).text = user_data["vesselName"] # vessel name

    doc.tables[0].cell(0, 0).text = "Vessel:"
    # doc.tables[0].cell(0, 0).width=cm(1.5)
    doc.tables[0].cell(0, 0).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0, 0).paragraphs[0].runs[0].font.name = 'Calibri' 
    doc.tables[0].cell(0, 0).paragraphs[0].runs[0].font.bold = True   

    doc.tables[0].cell(0, 1).text = data_pdf_para['Vessel']
    doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.bold = True

    # doc.tables[0].cell(0, 3).text = user_data["hullNumber"] # hull
    # doc.tables[0].cell(0, 2).width=Inches(0.5)
    doc.tables[0].cell(0, 2).text = "Hull no.:"
    doc.tables[0].cell(0, 2).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0, 2).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(0, 2).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(0, 3).text = data_pdf_para['Hull_No']
    doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.bold = True

    # doc.tables[0].cell(0, 5).text=user_data["vesselType"] # vessel type
    doc.tables[0].cell(0,4).text="Vessel Type:"
    doc.tables[0].cell(0,4).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0,4).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(0,4).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(0,5).text=data_pdf_para['Vessel_Type']
    doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(0,6).text="Local Start Time:"  #Local Start Time
    doc.tables[0].cell(0,6).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0,6).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(0,6).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(0,7).text=data_pdf_para['Local_Start_Time']  #Local Start Time
    doc.tables[0].cell(0,7).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(0,7).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[0].cell(0,7).paragraphs[0].runs[0].font.bold = True

    # # Table 1
    doc.tables[0].cell(1, 0).text = "Time Zone:"
    doc.tables[0].cell(1, 0).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(1, 0).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(1, 1).text = data_pdf_para['Time_Zone1'] 
    doc.tables[0].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[1].cell(0, 1).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(1, 2).text = "Local End Time:"
    doc.tables[0].cell(1, 2).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 2).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(1, 2).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(1, 3).text = data_pdf_para['Local_End_Time'] 
    doc.tables[0].cell(1, 3).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[1].cell(0, 3).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(1, 4).text = "Time Zone:"
    doc.tables[0].cell(1, 4).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 4).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(1, 4).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(1, 5).text = data_pdf_para['Time_Zone2'] 
    doc.tables[0].cell(1, 5).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 5).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[1].cell(0, 5).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(1, 6).text = "Form No.:" 
    doc.tables[0].cell(1, 6).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 6).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(1, 6).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(1, 7).text = data_pdf_para['Form_No'] 
    doc.tables[0].cell(1, 7).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(1, 7).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[1].cell(0, 7).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(2, 0).text = "IMO No."
    doc.tables[0].cell(2, 0).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(2, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(2, 0).paragraphs[0].runs[0].font.bold = True   
    doc.tables[0].cell(2, 1).text = data_pdf_para['IMO_No'] 
    doc.tables[0].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(2, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[1].cell(1, 1).paragraphs[0].runs[0].font.bold = True

    # # Table 2
    # Engine_info  =data_id1['Engine_info']
    # doc.tables[2].cell(0, 1).text = user_data["manufacture"] # make
    doc.tables[0].cell(4, 0).text = "Maker:"
    doc.tables[0].cell(4, 0).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(4, 0).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(4, 1).text = data_pdf_para['Maker']
    doc.tables[0].cell(4, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(4, 2).text = "Model:"
    doc.tables[0].cell(4, 2).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 2).paragraphs[0].runs[0].font.name = 'Calibri' 
    doc.tables[0].cell(4, 2).paragraphs[0].runs[0].font.bold = True   
    doc.tables[0].cell(4, 3).text = data_pdf_para['Model']
    doc.tables[0].cell(4, 3).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(0, 3).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(4, 4).text = "Licence:"
    doc.tables[0].cell(4, 4).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 4).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(4, 4).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(4, 5).text = data_pdf_para['License_Builder']
    doc.tables[0].cell(4, 5).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 5).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(0, 5).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(4, 6).text = "Serial No.:"
    doc.tables[0].cell(4, 6).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 6).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(4, 6).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(4, 7).text = data_pdf_para['Serial_No']
    doc.tables[0].cell(4, 7).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(4, 7).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(0, 7).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(5, 0).text = "MCR:"
    doc.tables[0].cell(5, 0).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(5, 0).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(5, 1).text = data_pdf_para['MCR']
    doc.tables[0].cell(5, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(1, 1).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(5, 2).text = "Speed at MCR:"
    doc.tables[0].cell(5, 2).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 2).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(5, 2).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(5, 3).text = data_pdf_para['Speed_at_MCR']
    doc.tables[0].cell(5, 3).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(1, 3).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(5, 4).text = "Bore:"
    doc.tables[0].cell(5, 4).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 4).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(5, 4).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(5, 5).text = data_pdf_para['Bore']
    doc.tables[0].cell(5, 5).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 5).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(1, 5).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(5, 6).text = "Stroke:"
    doc.tables[0].cell(5, 6).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 6).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(5, 6).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(5, 7).text = data_pdf_para['Stroke']
    doc.tables[0].cell(5, 7).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(5, 7).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[2].cell(1, 7).paragraphs[0].runs[0].font.bold = True

    ## Table3
    # Turbocharger_info  =data_id1['Turbocharger_info']
    doc.tables[0].cell(7, 0).text = "Maker"
    doc.tables[0].cell(7, 0).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(7, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(7, 0).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(7, 1).text = data_pdf_para['Maker_T']
    doc.tables[0].cell(7, 1).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(7, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[3].cell(0, 1).paragraphs[0].runs[0].font.bold = True

    doc.tables[0].cell(7, 2).text = "Model"
    doc.tables[0].cell(7, 2).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(7, 2).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[0].cell(7, 2).paragraphs[0].runs[0].font.bold = True
    doc.tables[0].cell(7, 3).text = data_pdf_para['Model_T']
    doc.tables[0].cell(7, 3).paragraphs[0].runs[0].font.size = Pt(7)
    doc.tables[0].cell(7, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    # doc.tables[3].cell(0, 3).paragraphs[0].runs[0].font.bold = True

    # General_Data  =data_id1['General_Data']
    # doc.tables[1].cell(1,1).paragraphs[0].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # doc.tables[1].cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.BOTTOM
    doc.tables[1].cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.tables[1].cell(1, 1).text = data_pdf_para['Total_Running_Hour'] #  Total running hours
    doc.tables[1].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(1, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    

    # doc.tables[1].cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.tables[1].cell(1, 3).text = data_pdf_para['Position'] #  Position
    doc.tables[1].cell(1, 3).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(1, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(1, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    doc.tables[1].cell(2, 1).text = data_pdf_para['Cylinder_Oil_type'] # Cyl oil type
    doc.tables[1].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(2, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    doc.tables[1].cell(2, 3).text = data_pdf_para['Cylinder_Oil_feed_rate'] # Cylinder Oil feed rate
    doc.tables[1].cell(2, 3).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(2, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(2, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

   

    doc.tables[1].cell(3,1).text = data_pdf_para['Normal_service_load_in_percentage_of_MCR'] # Normal service load in % of MCR
    doc.tables[1].cell(3,1).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(3,1).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    doc.tables[1].cell(3,3).text = data_pdf_para['Inspected_by_Rank'] # Inspected by (Rank)
    doc.tables[1].cell(3,3).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(3,3).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(3, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    
    doc.tables[1].cell(4, 1).text = data_pdf_para['Scrubber'] # Scrubber
    doc.tables[1].cell(4, 1).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(4, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(4, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    doc.tables[1].cell(4, 3).text = data_pdf_para['Fuel_Sulphur_percentage'] # Fuel Sulphur %
    doc.tables[1].cell(4, 3).paragraphs[0].runs[0].font.size = Pt(9)
    doc.tables[1].cell(4, 3).paragraphs[0].runs[0].font.name = 'Calibri'
    doc.tables[1].cell(4, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    
   
    list_ind=[1,2,3,4]
    doc.tables[3].style ='TableGrid' # Deposits section --tables[4]
    # print("selection_deposits=",selection_deposits)
    for i,k in zip(range(len(selection_deposits)),cyl_num):
        row_cells = doc.tables[3].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            # row_cells[j].text = 'l' if selection_deposits[i][ind]=='*' else selection_deposits[i][ind] 
            row_cells[j].text = 'l' if selection_deposits[i][ind]=='*' else 'LC'
            # if selection_deposits[i][ind]=='*':
            #     row_cells[j].text = 'l'
            # elif selection_deposits[i][ind]=='5':
            #     row_cells[j].text = 'LC'

           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        row_cells[5].text = 'l' 
        # row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[5].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[5].paragraphs[0].runs[0].font.name = 'Arial'
    #Mean row
    row_cells = doc.tables[3].add_row().cells
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
    # Breakage section --tables[4]
    list_ind=[1,2,3,-2]
    doc.tables[4].style ='TableGrid' 
    for i,k in zip(range(len(selection_brekage)),cyl_num):
        row_cells = doc.tables[4].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            # row_cells[j].text ='l' if selection_brekage[i][ind]=='*' else selection_brekage[i][ind]
            row_cells[j].text ='l' if selection_brekage[i][ind]=='*' else 'C'
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
        row_cells[-1].text = 'l' 
        # row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[-1].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[-1].paragraphs[0].runs[0].font.name = 'Arial'

    row_cells = doc.tables[4].add_row().cells
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Surface section --tables[7]
    list_ind=[1,2,3,4]
    list_ind_extra=[5,6,7,-2,-1]
    doc.tables[7].style ='TableGrid' 
    for i,k in zip(range(len(selection_surface)),cyl_num):
        row_cells = doc.tables[7].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            # row_cells[j].text = 'Cl' if selection_surface[i][ind]=='*' else selection_surface[i][ind]
            row_cells[j].text = 'Cl' if selection_surface[i][ind]=='*' else "S"
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        for ext in list_ind_extra:
            row_cells[ext].text = 'Cl' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[ext].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[ext].paragraphs[0].runs[0].font.name = 'Arial'
        
    row_cells = doc.tables[7].add_row().cells #Mean row
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Lubrication section --tables[8]
    list_ind=[1,2,3,4]
    list_ind_extra=[5,6,-2,-1]
    doc.tables[8].style ='TableGrid' 
    for i,k in zip(range(len(selection_lubrication)),cyl_num):
        row_cells = doc.tables[8].add_row().cells
        ind=0
        row_cells[0].text = str(k)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            # row_cells[j].text = 'N' if selection_lubrication[i][ind]=='*' else selection_lubrication[i][ind]
            row_cells[j].text = 'N' if selection_lubrication[i][ind]=='*' else 'O' if selection_lubrication[i][ind]=='2' else 'OB' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
        for ext in list_ind_extra: # Ring5 , Piston Skirt,Rod,Liner
            row_cells[ext].text = 'N' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[ext].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[ext].paragraphs[0].runs[0].font.name = 'Arial'

    row_cells = doc.tables[8].add_row().cells  #Mean row
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

########### adding images #########################
    
 
    table = doc.add_table(rows=math.ceil(len(img_list)/2), cols=2)
    table.style ='TableGrid' 

    im_count=0
    im_id=[]
    for i in cyl_num:
        im_id.append(int(i))  
    # print(img_list)

    for row,im_num in zip(table.rows,im_id):
        for cell in row.cells:
            im_count+=1
            
            if im_count<= len(img_list): # odd number of cylinders
                cell.text = "Cylinder "+str(im_id[im_count-1])
                p=cell.add_paragraph()
                r = p.add_run()
                r.add_picture(img_list[im_count-1],width=Inches(3))
    doc.save("files/report.docx")
    #below 2 lines are for linux
    # args = ["abiword", "--to", "report_output.pdf", "report.docx"  ]
    run(['abiword', '--to=pdf', 'files/report.docx'])
    # call(args )
    #below 2 lines are for windows
    # pythoncom.CoInitialize()
    # convert("files/report.docx") 

    # return send_file("./report.pdf", as_attachment=True,mimetype='application/pdf')
    
    # with open("files/report.pdf", "rb") as pdf_file: 
    #     encoded_string = base64.b64encode(pdf_file.read()) 
    data_pdf={'PDF':"https://scavaiapi.azurewebsites.net/files/report.pdf"} 
    # data_pdf['PDF']= [t.encode('utf-8') for t in title]

    for i,j in zip(data_k,obj_lst): ##Final Object
        cyl_obj[i]=j
    final_cyl_obj={**cyl_obj,**data_pdf, **data_id}
    # print("data_k:",data_k)
    # print("obj_lst:",obj_lst)
    # print("cyl_pred:",cyl_pred)
    # print("final_cyl_obj:",final_cyl_obj)
    # print("Object:", final_cyl_obj)
    result_obj = json.loads(json.dumps(final_cyl_obj))  
    # print(result_obj)
    return result_obj
